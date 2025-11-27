"""
Module để đọc các loại file khác nhau
Hiện tại hỗ trợ: DOCX, PDF, TXT
"""

import docx
import io
import re
from typing import Optional

def read_docx(file) -> Optional[str]:
    """
    Đọc nội dung từ file DOCX
    
    Args:
        file: File object từ Flask request
        
    Returns:
        str: Nội dung văn bản hoặc None nếu lỗi
    """
    try:
        # Đọc file từ memory
        doc = docx.Document(io.BytesIO(file.read()))
        
        # Trích xuất text từ các paragraph
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Bỏ qua paragraph trống
                full_text.append(paragraph.text.strip())
        
        # Trích xuất text từ tables (nếu có)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        # Ghép thành văn bản hoàn chỉnh
        content = "\n".join(full_text)
        
        # Làm sạch văn bản
        content = clean_text(content)
        
        return content if content.strip() else None
        
    except Exception as e:
        print(f"Lỗi khi đọc file DOCX: {str(e)}")
        return None

def read_txt(file) -> Optional[str]:
    """
    Đọc nội dung từ file TXT
    
    Args:
        file: File object từ Flask request
        
    Returns:
        str: Nội dung văn bản hoặc None nếu lỗi
    """
    try:
        # Thử các encoding phổ biến
        encodings = ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1']
        
        content = None
        file_bytes = file.read()
        
        for encoding in encodings:
            try:
                content = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return None
            
        # Làm sạch văn bản
        content = clean_text(content)
        
        return content if content.strip() else None
        
    except Exception as e:
        print(f"Lỗi khi đọc file TXT: {str(e)}")
        return None

def read_pdf(file) -> Optional[str]:
    """
    Đọc nội dung từ file PDF
    Requires: pip install PyPDF2
    
    Args:
        file: File object từ Flask request
        
    Returns:
        str: Nội dung văn bản hoặc None nếu lỗi
    """
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        
        full_text = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text.strip():
                full_text.append(text.strip())
        
        content = "\n".join(full_text)
        
        # Làm sạch văn bản
        content = clean_text(content)
        
        return content if content.strip() else None
        
    except ImportError:
        print("Cần cài đặt PyPDF2: pip install PyPDF2")
        return None
    except Exception as e:
        print(f"Lỗi khi đọc file PDF: {str(e)}")
        return None

def clean_text(text: str) -> str:
    """
    Làm sạch văn bản
    
    Args:
        text: Văn bản cần làm sạch
        
    Returns:
        str: Văn bản đã được làm sạch
    """
    if not text:
        return ""
    
    # Loại bỏ ký tự điều khiển
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
    
    # Chuẩn hóa khoảng trắng
    text = re.sub(r'\s+', ' ', text)
    
    # Loại bỏ dòng trống thừa
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Trim
    text = text.strip()
    
    return text

def get_file_type(filename: str) -> str:
    if not filename:
        return "unknown"
    
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.docx'):
        return "docx"
    elif filename_lower.endswith('.pdf'):
        return "pdf"
    elif filename_lower.endswith('.txt'):
        return "txt"
    else:
        return "unknown"

def read_file_auto(file) -> Optional[str]:
    """
    Read file
    
    Args:
        file: File object
        
    Returns:
        str: Content or error
    """
    file_type = get_file_type(file.filename)
    
    if file_type == "docx":
        return read_docx(file)
    elif file_type == "pdf":
        return read_pdf(file)
    elif file_type == "txt":
        return read_txt(file)
    else:
        return None